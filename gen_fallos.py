"""Generate Fallos_v4.docx — ELIMINATE respondents from full Supabase export."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import json, sys, os
from pathlib import Path
sys.path.insert(0, os.path.dirname(__file__))
import quality_check as qc

# Use the full export (all 1651 rows); fall back to older file if not found
full_path = Path(r'C:\Users\ajulve\Downloads\selmark_respuestas_full.json')
old_path  = Path(r'C:\Users\ajulve\Downloads\selmark_respuestas_2026-06-12 (1).json')
src = full_path if full_path.exists() else old_path
print(f"Loading from: {src}")
with open(src, encoding='utf-8') as f:
    data = json.load(f)
completed = [r for r in data if r.get('status') == 'complete']
results   = [qc.assess_response(r) for r in completed]

# Build human-readable reason from flag_summary
def human_reason(flag_summary):
    parts = []
    for segment in flag_summary.split('; '):
        if segment.startswith('short_time'):
            parts.append('Too fast (%s)' % segment.split(': ')[1].split(' <')[0])
        elif segment.startswith('straight_line_lifestyle'):
            detail = segment.split(': ', 1)[1] if ': ' in segment else ''
            parts.append('Straight-lining lifestyle (%s)' % detail[:60])
        elif segment.startswith('straight_line_attributes'):
            detail = segment.split(': ', 1)[1] if ': ' in segment else ''
            parts.append('Straight-lining attributes (%s)' % detail[:60])
        elif segment.startswith('copy_paste'):
            detail = segment.split(': ', 1)[1] if ': ' in segment else ''
            parts.append('Copy-paste (%s)' % detail[:60])
        elif segment.startswith('poor_open_text'):
            detail = segment.split(': ', 1)[1] if ': ' in segment else ''
            parts.append('Poor open text (%s)' % detail[:50])
        elif segment.startswith('short_why_answers'):
            detail = segment.split(': ', 1)[1] if ': ' in segment else ''
            parts.append('Non-answers in why questions (%s)' % detail[:70])
        elif segment.startswith('invalid_demographics'):
            detail = segment.split(': ', 1)[1] if ': ' in segment else ''
            parts.append('Invalid demographics (%s)' % detail[:50])
        elif segment.startswith('seesaw'):
            parts.append('Seesaw pattern in scale')
        elif segment.startswith('conditional_logic'):
            parts.append('Skip-logic violation')
    return '. '.join(parts) if parts else flag_summary[:120]

eliminate_rows = [
    (r['respondent_id'], r['duration_min'], human_reason(r['flag_summary']))
    for r in results if r['verdict'] == 'ELIMINATE'
]
eliminate_rows.sort(key=lambda x: float(x[1]) if x[1] != 'N/A' else 999)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_col_width(cell, width_cm):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.insert(0, tcW)

doc = Document()
section = doc.sections[0]
section.page_width    = Cm(29.7)
section.page_height   = Cm(21.0)
section.left_margin   = Cm(1.5)
section.right_margin  = Cm(1.5)
section.top_margin    = Cm(1.5)
section.bottom_margin = Cm(1.5)

title = doc.add_paragraph()
run = title.add_run('Selmark Survey QC — Respondents to Eliminate')
run.font.name = 'Arial'; run.font.size = Pt(14); run.font.bold = True
run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

sub = doc.add_paragraph()
n_ok   = sum(1 for r in results if r['verdict'] == 'OK')
n_rev  = sum(1 for r in results if r['verdict'] == 'REVIEW')
n_elim = len(eliminate_rows)
run2 = sub.add_run(
    f'Source: Supabase full export {src.stem}  |  Total responses: {len(completed)}  |  '
    f'OK: {n_ok}  REVIEW: {n_rev}  ELIMINATE: {n_elim}'
)
run2.font.name = 'Arial'; run2.font.size = Pt(9)
run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

table = doc.add_table(rows=1, cols=3)
table.style = 'Table Grid'
WIDTHS = [8.5, 1.2, 17.0]

hdr = table.rows[0]
for i, (cell, text) in enumerate(zip(hdr.cells, ['Respondent ID', 'Min', 'Reason for elimination'])):
    cell.text = text
    set_cell_bg(cell, '1F3864')
    set_col_width(cell, WIDTHS[i])
    for para in cell.paragraphs:
        for run in para.runs:
            run.font.name = 'Arial'; run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for idx, (rid, dur, reason) in enumerate(eliminate_rows):
    row = table.add_row()
    shade = 'FFFFFF' if idx % 2 == 0 else 'F2F2F2'
    for i, (cell, text) in enumerate(zip(row.cells, [rid, dur, reason])):
        cell.text = text
        set_cell_bg(cell, shade)
        set_col_width(cell, WIDTHS[i])
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.name = 'Courier New' if i == 0 else 'Arial'
                run.font.size = Pt(8) if i == 0 else Pt(9)

out = r'C:\Users\ajulve\Downloads\Fallos_v5.docx'
doc.save(out)
print(f'Saved: {out}  ({n_elim} rows  |  OK={n_ok}  REVIEW={n_rev}  ELIMINATE={n_elim})')

# ── Generate Excel for Cint ──────────────────────────────────────────────────
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

wb = openpyxl.Workbook()
ws = wb.active
ws.title = 'Disqualified Respondents'

header_fill = PatternFill('solid', fgColor='1F3864')
header_font = Font(name='Arial', bold=True, color='FFFFFF', size=10)
for col, text in enumerate(['Respondent ID', 'Duration (min)', 'Reason'], start=1):
    cell = ws.cell(row=1, column=col, value=text)
    cell.font = header_font
    cell.fill = header_fill
    cell.alignment = Alignment(horizontal='center')

row_fills = [PatternFill('solid', fgColor='FFFFFF'), PatternFill('solid', fgColor='F2F2F2')]
for idx, (rid, dur, reason) in enumerate(eliminate_rows, start=2):
    fill = row_fills[idx % 2]
    for col, val in enumerate([rid, dur, reason], start=1):
        cell = ws.cell(row=idx, column=col, value=val)
        cell.font = Font(name='Arial', size=9)
        cell.fill = fill

ws.column_dimensions['A'].width = 40
ws.column_dimensions['B'].width = 14
ws.column_dimensions['C'].width = 80

xl_out = r'C:\Users\ajulve\Downloads\Selmark_Disqualified_IDs.xlsx'
wb.save(xl_out)
print(f'Saved Excel: {xl_out}  ({n_elim} rows)')
